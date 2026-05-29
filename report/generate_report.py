from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

title = doc.add_heading('Job Market Analytics Pipeline Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Assignment #3 - Tools & Techniques for Data Science')
doc.add_paragraph('University of Central Punjab')
doc.add_paragraph('Faculty of IT & CS - Department of Applied Computing & Technologies')

doc.add_page_break()

doc.add_heading('1. Introduction and Problem Statement', level=1)
doc.add_paragraph(
    'This report documents the development and implementation of an automated job market analytics pipeline designed to collect, clean, and analyze AI, ML, and Data-related job postings from multiple online sources. '
    'The project addresses the challenge of monitoring job markets manually, which is slow and inconsistent due to different data formats across websites.'
)
doc.add_paragraph(
    'As a junior data engineer in a recruitment analytics company, I built an end-to-end pipeline that extracts job data from free public APIs, standardizes the schema, cleans the data using KNIME, orchestrates the workflow with Apache Airflow, and sends automated notifications through n8n.'
)

doc.add_heading('2. Source Selection and API Links', level=1)

doc.add_heading('2.1 Data Sources Used', level=2)
sources_table = doc.add_table(rows=5, cols=4)
sources_table.style = 'Table Grid'
headers = ['Source', 'Purpose', 'API Endpoint', 'Notes']
for i, header in enumerate(headers):
    sources_table.rows[0].cells[i].text = header

sources_data = [
    ['Arbeitnow', 'General jobs with Europe focus', 'https://www.arbeitnow.com/api/job-board-api', 'Free, no API key'],
    ['RemoteOK', 'Remote job listings with salary', 'https://remoteok.com/api', 'Public JSON feed'],
    ['Himalayas', 'Remote jobs with filters', 'https://himalayas.app/jobs/api/search?q=data', 'Free public JSON API'],
    ['RemoteJobs.org', 'Remote jobs by category', 'https://remotejobs.org/api/v1/jobs', 'Free, no signup']
]
for i, row_data in enumerate(sources_data):
    for j, cell_data in enumerate(row_data):
        sources_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('2.2 Schema Mapping', level=2)
doc.add_paragraph(
    'The four sources return different raw schemas. I mapped all sources to a standardized schema with fields: '
    'source, job_id, title, company_name, location_raw, remote_status, job_type, category_raw, tags_raw, description, '
    'publication_date, job_url, salary_text_raw, salary_min_raw, salary_max_raw, currency_raw, salary_min_usd, '
    'salary_max_usd, salary_mid_usd, experience_years_min, experience_years_max, experience_bracket, extracted_skills, '
    'job_category_clean, scrape_date.'
)

doc.add_heading('3. Pipeline Architecture', level=1)
doc.add_paragraph('The end-to-end pipeline follows this architecture:')
doc.add_paragraph('1. Extraction: Four Python scripts collect raw data from each API')
doc.add_paragraph('2. Merging: A merge script standardizes schemas and combines all sources')
doc.add_paragraph('3. Cleaning: KNIME workflow cleans, filters, and transforms the data')
doc.add_paragraph('4. Validation: Python script performs data quality checks')
doc.add_paragraph('5. Metrics: Calculate analysis metrics and generate summary')
doc.add_paragraph('6. Notification: n8n workflow sends email/webhook notification')
doc.add_paragraph('7. Orchestration: Apache Airflow DAG orchestrates all tasks')

doc.add_paragraph('Pipeline Architecture:')
try:
    doc.add_picture('C:\\Users\\Tahan\\Desktop\\Assignment 3\\screenshots\\airflow\\Screenshot 2026-05-19 013341.png', width=Inches(6.5))
except:
    doc.add_paragraph('[Airflow DAG Screenshot]')

doc.add_heading('4. Extraction Process', level=1)
doc.add_paragraph('Four separate Python scripts were created to extract data from each source:')
doc.add_paragraph('- extract_arbeitnow.py: Fetches jobs from Arbeitnow API')
doc.add_paragraph('- extract_remoteok.py: Fetches jobs from RemoteOK API')
doc.add_paragraph('- extract_himalayas.py: Fetches jobs from Himalayas API')
doc.add_paragraph('- extract_remotejobs.py: Fetches jobs from RemoteJobs.org API')

doc.add_heading('4.1 Raw Data Statistics', level=2)
raw_stats_table = doc.add_table(rows=5, cols=3)
raw_stats_table.style = 'Table Grid'
raw_stats_table.rows[0].cells[0].text = 'Source'
raw_stats_table.rows[0].cells[1].text = 'Jobs Collected'
raw_stats_table.rows[0].cells[2].text = 'Percentage'

sources_raw = [
    ['Himalayas', '118', '29.4%'],
    ['Arbeitnow', '100', '24.9%'],
    ['RemoteOK', '99', '24.6%'],
    ['RemoteJobs.org', '85', '21.1%']
]
for i, row_data in enumerate(sources_raw):
    for j, cell_data in enumerate(row_data):
        raw_stats_table.rows[i+1].cells[j].text = cell_data

doc.add_paragraph('Total raw jobs collected: 402')

doc.add_heading('5. KNIME Cleaning and Transformation', level=1)
doc.add_paragraph('KNIME workflow performs the following cleaning operations:')
doc.add_paragraph('- Remove HTML tags and special characters from descriptions')
doc.add_paragraph('- Standardize title, company name, and location columns')
doc.add_paragraph('- Deduplicate jobs using job_url')
doc.add_paragraph('- Apply AI/ML/Data relevance filter')
doc.add_paragraph('- Extract and standardize skills from title, tags, and description')
doc.add_paragraph('- Create job_category_clean column')
doc.add_paragraph('- Standardize remote_status (Remote, On-site, Hybrid, Unknown)')
doc.add_paragraph('- Extract experience years and create experience_bracket')
doc.add_paragraph('- Extract salary, normalize pay period, and convert currencies to USD')

try:
    doc.add_picture('C:\\Users\\Tahan\\Desktop\\Assignment 3\\screenshots\\knime\\Screenshot 2026-05-19 013821.png', width=Inches(6.5))
except:
    doc.add_paragraph('[KNIME Workflow Screenshot]')

doc.add_heading('6. AI/ML/Data Filtering', level=1)
doc.add_paragraph('Filtering logic checks job titles and descriptions for relevant keywords:')
doc.add_paragraph('- Data Analytics: data analyst, reporting analyst, product analyst')
doc.add_paragraph('- Data Science: data scientist, ML, NLP, predictive modeling')
doc.add_paragraph('- Data Engineering: data engineer, ETL, pipeline, warehouse')
doc.add_paragraph('- AI/ML: machine learning, AI engineer, deep learning, LLM')
doc.add_paragraph('- Business Intelligence: BI analyst, Power BI, Tableau')
doc.add_paragraph('- Analytics Engineering: analytics engineer, dbt, metrics')

doc.add_heading('6.1 Filtering Results', level=2)
doc.add_paragraph('Jobs before filtering: 402')
doc.add_paragraph('Jobs after filtering: 223')
doc.add_paragraph('Filter drop rate: 44.53%')

doc.add_heading('7. Salary Extraction and USD Conversion', level=1)
doc.add_paragraph('Salary data was extracted from structured salary fields and description text.')
doc.add_paragraph('Currency conversion rates used (as of 2026-05-18):')
doc.add_paragraph('- USD: 1.00')
doc.add_paragraph('- EUR: 1.08')
doc.add_paragraph('- GBP: 1.25')
doc.add_paragraph('- PKR: 0.0036')

doc.add_paragraph('Hourly and monthly salaries were converted to annual using:')
doc.add_paragraph('- Hourly: 40 hours/week x 52 weeks = 2,080 hours/year')
doc.add_paragraph('- Monthly: x 12 months/year')

doc.add_heading('8. Experience Year Bracket Extraction', level=1)
doc.add_paragraph('Experience requirements were extracted from job titles and descriptions:')
doc.add_paragraph('- 0-1 years: Entry level, fresh graduate, internship, trainee')
doc.add_paragraph('- 1-3 years: Junior/mid, 1+ years, 2 years')
doc.add_paragraph('- 3-5 years: Mid-level, 3+ years, 4 years')
doc.add_paragraph('- 5-8 years: Senior, 5+ years, 6-8 years')
doc.add_paragraph('- 8+ years: Principal, staff, lead, director')
doc.add_paragraph('- Not mentioned: No clear experience requirement')

doc.add_heading('9. Airflow DAG Explanation', level=1)
doc.add_paragraph('The Airflow DAG orchestrates the complete pipeline with the following tasks:')
doc.add_paragraph('1. extract_arbeitnow: Collect raw data from Arbeitnow (parallel)')
doc.add_paragraph('2. extract_remoteok: Collect raw data from RemoteOK (parallel)')
doc.add_paragraph('3. extract_himalayas: Collect raw data from Himalayas (parallel)')
doc.add_paragraph('4. extract_remotejobs: Collect raw data from RemoteJobs.org (parallel)')
doc.add_paragraph('5. merge_sources: Standardize schemas and merge all datasets')
doc.add_paragraph('6. run_knime_workflow: Execute KNIME cleaning workflow')
doc.add_paragraph('7. validate_clean_output: Check row counts, columns, duplicates')
doc.add_paragraph('8. calculate_metrics: Generate analysis metrics')
doc.add_paragraph('9. trigger_n8n_workflow: Send notification to n8n webhook')
doc.add_paragraph('10. archive_outputs: Save timestamped copies of output files')

try:
    doc.add_picture('C:\\Users\\Tahan\\Desktop\\Assignment 3\\screenshots\\airflow\\Screenshot 2026-05-19 013341.png', width=Inches(6.5))
except:
    doc.add_paragraph('[Airflow DAG Screenshot]')

doc.add_heading('10. n8n Workflow Explanation', level=1)
doc.add_paragraph('n8n workflow is triggered via HTTP webhook after metrics are calculated.')
doc.add_paragraph('The workflow includes:')
doc.add_paragraph('- HTTP Request node to receive pipeline status')
doc.add_paragraph('- Email notification with job summary')
doc.add_paragraph('- Summary includes: total jobs, jobs by source, remote ratio, 0-1 year jobs, average salary, pipeline status')

try:
    doc.add_picture('C:\\Users\\Tahan\\Desktop\\Assignment 3\\screenshots\\n8n\\Screenshot 2026-05-19 013419.png', width=Inches(6.5))
except:
    doc.add_paragraph('[n8n Workflow Screenshot]')

doc.add_heading('11. Data Quality Checks', level=1)

doc.add_heading('11.1 Quality Check Results', level=2)
quality_table = doc.add_table(rows=9, cols=3)
quality_table.style = 'Table Grid'
quality_table.rows[0].cells[0].text = 'Check'
quality_table.rows[0].cells[1].text = 'Result'
quality_table.rows[0].cells[2].text = 'Status'

quality_checks = [
    ['API Response', 'All 4 sources returned successful responses', 'Pass'],
    ['Source Count', '402 total raw jobs from 4 sources', 'Pass'],
    ['Schema', 'All 24 mandatory columns present', 'Pass'],
    ['Duplicates', 'Duplicates removed during cleaning', 'Pass'],
    ['Relevance Filter', '402 -> 223 after filtering', 'Pass'],
    ['Missing Values', '0% missing title/company', 'Pass'],
    ['Date Check', 'All publication dates valid', 'Pass'],
    ['Salary Check', 'Zeros converted to null', 'Pass']
]
for i, row_data in enumerate(quality_checks):
    for j, cell_data in enumerate(row_data):
        quality_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('12. Analysis Answers and Visualizations', level=1)

doc.add_heading('12.1 Jobs by Source', level=2)
source_table = doc.add_table(rows=5, cols=4)
source_table.style = 'Table Grid'
source_table.rows[0].cells[0].text = 'Source'
source_table.rows[0].cells[1].text = 'Before Filter'
source_table.rows[0].cells[2].text = 'After Filter'
source_table.rows[0].cells[3].text = 'Filter Rate'

source_analysis = [
    ['Himalayas', '118', '117', '0.8%'],
    ['RemoteJobs.org', '85', '50', '41.2%'],
    ['Arbeitnow', '100', '28', '72.0%'],
    ['RemoteOK', '99', '28', '71.7%']
]
for i, row_data in enumerate(source_analysis):
    for j, cell_data in enumerate(row_data):
        source_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('12.2 Remote/On-site Distribution', level=2)
doc.add_paragraph('Remote: 201 jobs (90.13%)')
doc.add_paragraph('On-site: 22 jobs (9.87%)')
doc.add_paragraph('Hybrid: 0 jobs (0%)')
doc.add_paragraph('Unknown: 0 jobs (0%)')

doc.add_heading('12.3 Experience Bracket Distribution', level=2)
exp_table = doc.add_table(rows=7, cols=3)
exp_table.style = 'Table Grid'
exp_table.rows[0].cells[0].text = 'Bracket'
exp_table.rows[0].cells[1].text = 'Count'
exp_table.rows[0].cells[2].text = 'Percentage'

exp_data = [
    ['Not mentioned', '74', '33.2%'],
    ['8+ years', '61', '27.4%'],
    ['5-8 years', '44', '19.7%'],
    ['0-1 years (Entry)', '32', '14.3%'],
    ['3-5 years', '8', '3.6%'],
    ['1-3 years', '4', '1.8%']
]
for i, row_data in enumerate(exp_data):
    for j, cell_data in enumerate(row_data):
        exp_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('12.4 Salary Analysis', level=2)
doc.add_paragraph('Overall Average Salary: $150,672.70 USD')
doc.add_paragraph('Median Salary: $136,070.00 USD')
doc.add_paragraph('Salary Data Coverage: 32.7% (73 out of 223 jobs)')

doc.add_heading('12.5 Average Salary by Category', level=3)
salary_cat_table = doc.add_table(rows=6, cols=2)
salary_cat_table.style = 'Table Grid'
salary_cat_table.rows[0].cells[0].text = 'Category'
salary_cat_table.rows[0].cells[1].text = 'Avg Salary (USD)'

salary_cat_data = [
    ['AI/ML', '$179,251.81'],
    ['Data Engineering', '$173,001.54'],
    ['Data Analytics', '$141,878.59'],
    ['Other Data', '$112,201.50'],
    ['Business Intelligence', '$78,050.00']
]
for i, row_data in enumerate(salary_cat_data):
    for j, cell_data in enumerate(row_data):
        salary_cat_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('12.6 Top Skills', level=2)
skills_table = doc.add_table(rows=11, cols=2)
skills_table.style = 'Table Grid'
skills_table.rows[0].cells[0].text = 'Skill'
skills_table.rows[0].cells[1].text = 'Count'

skills_data = [
    ['Scala', '39'],
    ['Machine Learning', '38'],
    ['Python', '23'],
    ['SQL', '23'],
    ['NLP', '16'],
    ['Excel', '11'],
    ['dbt', '10'],
    ['Azure', '8'],
    ['GCP', '8'],
    ['AWS', '7']
]
for i, row_data in enumerate(skills_data):
    for j, cell_data in enumerate(row_data):
        skills_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('12.7 Top Companies by Job Count', level=2)
companies_table = doc.add_table(rows=11, cols=2)
companies_table.style = 'Table Grid'
companies_table.rows[0].cells[0].text = 'Company'
companies_table.rows[0].cells[1].text = 'Job Count'

companies_data = [
    ['Mindrift', '7'],
    ['Mercor', '5'],
    ['NTT DATA', '4'],
    ['acemate.ai', '3'],
    ['Auralis Group', '3'],
    ['RWS Group', '3'],
    ['CI&T', '3'],
    ['Newfold Digital', '2'],
    ['Globaldev Group', '2'],
    ['Highmark Health', '2']
]
for i, row_data in enumerate(companies_data):
    for j, cell_data in enumerate(row_data):
        companies_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('12.8 Jobs by Category', level=2)
cat_table = doc.add_table(rows=6, cols=2)
cat_table.style = 'Table Grid'
cat_table.rows[0].cells[0].text = 'Category'
cat_table.rows[0].cells[1].text = 'Job Count'

cat_data = [
    ['Data Analytics', '75'],
    ['Other Data', '66'],
    ['Data Engineering', '49'],
    ['AI/ML', '31'],
    ['Business Intelligence', '2']
]
for i, row_data in enumerate(cat_data):
    for j, cell_data in enumerate(row_data):
        cat_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('12.9 Salary Coverage by Source', level=2)
coverage_table = doc.add_table(rows=5, cols=3)
coverage_table.style = 'Table Grid'
coverage_table.rows[0].cells[0].text = 'Source'
coverage_table.rows[0].cells[1].text = 'Coverage %'
coverage_table.rows[0].cells[2].text = 'Jobs with Salary'

coverage_data = [
    ['Himalayas', '43.6%', '51'],
    ['RemoteJobs.org', '36.0%', '18'],
    ['RemoteOK', '14.3%', '4'],
    ['Arbeitnow', '0.0%', '0']
]
for i, row_data in enumerate(coverage_data):
    for j, cell_data in enumerate(row_data):
        coverage_table.rows[i+1].cells[j].text = cell_data

doc.add_heading('13. Challenges and Solutions', level=1)
doc.add_paragraph('1. Schema Inconsistency: Different APIs used different field names - Solved by creating a standardized schema mapping.')
doc.add_paragraph('2. Salary Format Diversity: Salaries came in various formats - Solved by implementing conversion logic.')
doc.add_paragraph('3. Experience Extraction: No structured experience fields - Solved by keyword matching in titles and descriptions.')
doc.add_paragraph('4. Remote-only Sources: RemoteOK and Himalayas only list remote jobs - Documented this limitation in analysis.')

doc.add_heading('14. Conclusion', level=1)
doc.add_paragraph(
    'This project successfully demonstrates the implementation of a complete data pipeline for job market analytics. '
    'The pipeline collects data from four free job APIs, standardizes the schema, cleans and filters for AI/ML/Data jobs using KNIME, '
    'orchestrates the workflow with Apache Airflow, and sends automated notifications through n8n.'
)
doc.add_paragraph(
    'Key Results: 223 AI/ML/Data jobs collected, 90.1% remote positions, 32 entry-level (0-1 year) positions available, '
    'average salary of $150,672 USD, and Data Analytics as the top job category.'
)
doc.add_paragraph(
    'The pipeline is fully automated and can be scheduled to run periodically to keep track of the evolving job market.'
)

doc.add_heading('15. Appendix - Screenshots', level=1)

doc.add_paragraph('Airflow DAG:')
try:
    doc.add_picture('C:\\Users\\Tahan\\Desktop\\Assignment 3\\screenshots\\airflow\\Screenshot 2026-05-19 013341.png', width=Inches(6.5))
except:
    doc.add_paragraph('[Airflow DAG Screenshot - See screenshots/airflow folder]')

doc.add_paragraph('KNIME Workflow:')
try:
    doc.add_picture('C:\\Users\\Tahan\\Desktop\\Assignment 3\\screenshots\\knime\\Screenshot 2026-05-19 013821.png', width=Inches(6.5))
except:
    doc.add_paragraph('[KNIME Workflow Screenshot - See screenshots/knime folder]')

doc.add_paragraph('n8n Workflow:')
try:
    doc.add_picture('C:\\Users\\Tahan\\Desktop\\Assignment 3\\screenshots\\n8n\\Screenshot 2026-05-19 013419.png', width=Inches(6.5))
except:
    doc.add_paragraph('[n8n Workflow Screenshot - See screenshots/n8n folder]')

doc.add_paragraph('Cleaned Data Output:')
try:
    doc.add_picture('C:\\Users\\Tahan\\Desktop\\Assignment 3\\screenshots\\output\\Screenshot 2026-05-19 013509.png', width=Inches(6.5))
except:
    doc.add_paragraph('[Cleaned Data Output Screenshot - See screenshots/output folder]')

doc.save('C:\\Users\\Tahan\\Desktop\\Assignment 3\\report\\final_report.docx')
print("Report created: C:\\Users\\Tahan\\Desktop\\Assignment 3\\report\\final_report.docx")